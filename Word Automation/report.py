#!/usr/bin/env python
# coding: utf-8

# In[23]:


import pandas as pd
pd.set_option('display.float_format', lambda x: '%.1f' % x)
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
import os,sys


# In[2]:


cases_population=pd.read_csv('cases_population.csv')


# In[4]:


options=list(cases_population['location'])


# In[12]:


print ('List of availiable input options')
print("\n")
print(options)
print("\n")
country = input('Enter your country:')
print("\n")
print("Done! Your report is created in your folder named : {}_report.docx".format(country))


# In[14]:


selection = cases_population.loc[cases_population['location'] == country]


# In[15]:


blue_bar = (selection['total_cases'])
# Specify the values of orange bars (height)
orange_bar = (selection['population'])
country=list(selection['location'])
country=country[0]
total_cases=list(selection['total_cases'])
total_cases=total_cases[0]
population=list(selection['population'])
population=population[0]


# In[18]:


plt.figure(figsize=[6,4])
N = 1
# Width of a bar 
width = 0.3       
ind = np.arange(N)
# Plotting
plt.bar(ind, blue_bar , width, label='Cases')
plt.bar(ind + width, orange_bar, width, label='Population')

# plt.xlabel('Here goes x-axis label')
# plt.ylabel('Here goes y-axis label')
plt.title('{} Cases/Population bar chart'.format(country))
plt.legend(loc='best')
plt.savefig('graph.png')


# In[17]:


document = Document()
document.add_heading('{} cases/population'.format(country), level=0)
table = document.add_table(rows=2, cols=3)
table.cell(0, 0).text='LOCATION'
table.cell(0, 1).text='TOTAL CASES'
table.cell(0, 2).text='POPULATION'
table.cell(1, 0).text=str(country)
table.cell(1, 1).text=str(total_cases)
table.cell(1, 2).text=str(population)
document.add_picture('graph.png')
document.save('{}_report.docx'.format(country))


# In[44]:


folder_path = os.getcwd()
test = os.listdir(folder_path)


# In[45]:


for images in test:
    if images.endswith(".png"):
        os.remove(os.path.join(folder_path, images))


# In[ ]:




